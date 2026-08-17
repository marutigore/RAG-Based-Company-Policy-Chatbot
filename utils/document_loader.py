"""
Document Loader Module.
Extracts clean text and metadata from multiple corporate formats:
PDF, DOCX, XLSX, CSV, TXT, Markdown, and HTML.
"""

import os
import re
import csv
import zipfile
import logging
import xml.etree.ElementTree as ET
from typing import List, Dict, Any, Optional
import pymupdf as fitz  # PyMuPDF

# Initialize module logger
logger = logging.getLogger("utils.document_loader")


def clean_text(text: str) -> str:
    """
    Cleans extracted text by normalizing whitespace, resolving line wraps,
    and removing non-printable control characters.
    """
    cleaned = text.replace("\r", "\n")
    cleaned = cleaned.replace("\u2011", "-").replace("\u2010", "-").replace("\u2013", "-").replace("\u2014", "-")
    cleaned = cleaned.replace("\u201c", '"').replace("\u201d", '"').replace("\u2018", "'").replace("\u2019", "'")
    cleaned = re.sub(r"(\w+)-\n(\w+)", r"\1\2", cleaned)
    cleaned = re.sub(r"\n+", "\n", cleaned)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    return cleaned.strip()


def load_docx(file_path: str, filename: str) -> List[Dict[str, Any]]:
    """Loads DOCX files using python-docx or fallback built-in zipfile XML extraction."""
    paragraphs = []
    try:
        import docx
        doc = docx.Document(file_path)
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                paragraphs.append(txt)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(f"[Table Row]: {row_text}")
    except Exception:
        # Fallback to direct OOXML parsing
        try:
            with zipfile.ZipFile(file_path) as z:
                xml_content = z.read("word/document.xml")
                tree = ET.fromstring(xml_content)
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                for p in tree.iterfind('.//w:p', namespaces):
                    texts = [node.text for node in p.iterfind('.//w:t', namespaces) if node.text]
                    if texts:
                        paragraphs.append("".join(texts).strip())
        except Exception as zip_err:
            logger.error(f"DOCX extraction fallback failed: {zip_err}")
            raise ValueError(f"Failed to read DOCX file: {zip_err}")

    if not paragraphs:
        return []

    # Partition long DOCX into virtual pages of ~500 words
    full_text = "\n\n".join(paragraphs)
    words = full_text.split()
    page_size = 350
    pages_data = []
    
    for i in range(0, max(1, len(words)), page_size):
        chunk_words = words[i:i + page_size]
        page_num = (i // page_size) + 1
        page_text = clean_text(" ".join(chunk_words))
        if page_text:
            pages_data.append({
                "text": page_text,
                "metadata": {
                    "source": filename,
                    "page": page_num,
                    "format": "docx"
                }
            })
    return pages_data


def load_xlsx_or_csv(file_path: str, filename: str) -> List[Dict[str, Any]]:
    """Loads CSV and XLSX policy matrices and tabular records."""
    pages_data = []
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".csv":
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                reader = csv.reader(f)
                headers = next(reader, [])
                rows_text = []
                for idx, row in enumerate(reader):
                    row_repr = ", ".join(f"{h}: {val}" for h, val in zip(headers, row) if val)
                    if row_repr:
                        rows_text.append(f"Row {idx + 1}: {row_repr}")
                    if len(rows_text) >= 25:
                        pages_data.append({
                            "text": "\n".join(rows_text),
                            "metadata": {
                                "source": filename,
                                "page": len(pages_data) + 1,
                                "format": "csv"
                            }
                        })
                        rows_text = []
                if rows_text:
                    pages_data.append({
                        "text": "\n".join(rows_text),
                        "metadata": {
                            "source": filename,
                            "page": len(pages_data) + 1,
                            "format": "csv"
                        }
                    })
        except Exception as e:
            logger.error(f"CSV read error: {e}")
            raise ValueError(f"Failed to read CSV: {e}")
    else:
        # XLSX
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_lines = [f"[Sheet: {sheet_name}]"]
                for row in sheet.iter_rows(values_only=True):
                    vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
                    if vals:
                        sheet_lines.append(" | ".join(vals))
                if len(sheet_lines) > 1:
                    pages_data.append({
                        "text": clean_text("\n".join(sheet_lines)),
                        "metadata": {
                            "source": filename,
                            "page": len(pages_data) + 1,
                            "sheet": sheet_name,
                            "format": "xlsx"
                        }
                    })
        except Exception as e:
            logger.warning(f"openpyxl failed, trying fallback: {e}")
            raise ValueError(f"Failed to read XLSX: {e}")

    return pages_data


def load_text_file(file_path: str, filename: str) -> List[Dict[str, Any]]:
    """Loads plain text (.txt) and Markdown (.md) documents."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
    except Exception as e:
        logger.error(f"Error reading text file: {e}")
        raise ValueError(f"Failed to read text file: {e}")

    # Split into logical sections or pages (~1500 chars)
    sections = re.split(r'\n(?=#{1,3}\s)|\n\n---+\n\n', content)
    pages_data = []
    
    for idx, sec in enumerate(sections):
        cleaned = clean_text(sec)
        if cleaned:
            pages_data.append({
                "text": cleaned,
                "metadata": {
                    "source": filename,
                    "page": idx + 1,
                    "format": "text/markdown"
                }
            })
    
    if not pages_data and content.strip():
        pages_data.append({
            "text": clean_text(content),
            "metadata": {
                "source": filename,
                "page": 1,
                "format": "text"
            }
        })
    return pages_data


def load_html(file_path: str, filename: str) -> List[Dict[str, Any]]:
    """Loads HTML documents stripping tags while preserving textual structure."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            html_raw = f.read()
    except Exception as e:
        logger.error(f"HTML read error: {e}")
        raise ValueError(f"Failed to read HTML: {e}")

    # Strip scripts & styles
    clean_html = re.sub(r'<(script|style)[^>]*>.*?</\1>', '', html_raw, flags=re.DOTALL | re.IGNORECASE)
    # Convert paragraph / breaks to newlines
    clean_html = re.sub(r'<(p|br|div|tr|h[1-6])[^>]*>', '\n', clean_html, flags=re.IGNORECASE)
    # Strip all remaining tags
    plain_text = re.sub(r'<[^>]+>', '', clean_html)
    cleaned = clean_text(plain_text)
    
    return [{
        "text": cleaned,
        "metadata": {
            "source": filename,
            "page": 1,
            "format": "html"
        }
    }] if cleaned else []


def load_pdf(file_path: str, custom_filename: Optional[str] = None) -> List[Dict[str, Any]]:
    """
    Loads a PDF document page-by-page and extracts text and metadata.
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        raise FileNotFoundError(f"The file '{file_path}' does not exist.")

    pages_data: List[Dict[str, Any]] = []
    raw_filename = custom_filename or os.path.basename(file_path)
    filename = raw_filename.replace('\u2011', '-').encode('ascii', 'replace').decode('ascii').replace('?', '_')

    try:
        logger.info(f"Opening PDF document: {file_path}")
        doc = fitz.open(file_path)
        
        import config
        if len(doc) > config.MAX_PAGES_LIMIT:
            logger.error(f"Document {filename} has {len(doc)} pages, exceeding limit of {config.MAX_PAGES_LIMIT}.")
            doc.close()
            raise ValueError(f"Document exceeds maximum limit of {config.MAX_PAGES_LIMIT} pages.")
        
        for page_num in range(len(doc)):
            try:
                page = doc.load_page(page_num)
                raw_text = page.get_text()
                cleaned_page_text = clean_text(raw_text)
                
                if not cleaned_page_text:
                    continue
                
                page_info = {
                    "text": cleaned_page_text,
                    "metadata": {
                        "source": filename,
                        "page": page_num + 1,
                        "total_pages": len(doc),
                        "format": "pdf"
                    }
                }
                pages_data.append(page_info)
            except Exception as page_err:
                logger.error(f"Failed to process page {page_num + 1} in {filename}: {page_err}")
                continue
            
        logger.info(f"Successfully loaded {len(pages_data)} pages from {filename}")
        doc.close()
        
    except Exception as e:
        logger.error(f"Error loading PDF {file_path}: {e}")
        raise Exception(f"Failed to read PDF: {e}")

    return pages_data


def load_document(file_path: str, custom_filename: Optional[str] = None) -> List[Dict[str, Any]]:
    """
    Central multi-format dispatcher.
    Parses PDF, DOCX, XLSX, CSV, TXT, MD, and HTML files.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File '{file_path}' does not exist.")

    raw_filename = custom_filename or os.path.basename(file_path)
    filename = raw_filename.replace('\u2011', '-').encode('ascii', 'replace').decode('ascii').replace('?', '_')
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return load_pdf(file_path, custom_filename=filename)
    elif ext in [".docx", ".doc"]:
        return load_docx(file_path, filename=filename)
    elif ext in [".xlsx", ".csv"]:
        return load_xlsx_or_csv(file_path, filename=filename)
    elif ext in [".txt", ".md", ".markdown"]:
        return load_text_file(file_path, filename=filename)
    elif ext in [".html", ".htm"]:
        return load_html(file_path, filename=filename)
    else:
        # Fallback to plain text reader
        return load_text_file(file_path, filename=filename)
